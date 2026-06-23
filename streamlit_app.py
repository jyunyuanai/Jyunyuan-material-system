from __future__ import annotations

import base64
import getpass
import hashlib
import hmac
import re
import secrets
import sys
import time
from copy import deepcopy
from dataclasses import dataclass
from enum import StrEnum
from io import BytesIO
from pathlib import PurePosixPath
from typing import Any, Mapping
from zipfile import ZIP_DEFLATED, BadZipFile, ZipFile

import streamlit as st
from docx import Document
from docxcompose.composer import Composer
from lxml import etree


# =============================================================================
# 程式設定與共用資料
# =============================================================================

APP_NAME = "監造計畫書材料系統"
MAX_LOGIN_FAILURES = 5
LOGIN_COOLDOWN_SECONDS = 30
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
NS = {"w": W_NS}
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

MAX_DOCX_BYTES = 25 * 1024 * 1024
MAX_ZIP_ENTRIES = 500
MAX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
PROHIBITED_PART_PREFIXES = ("word/embeddings/", "word/activeX/")
PROHIBITED_PART_NAMES = {"word/vbaProject.bin"}

SCRYPT_N = 2**14
SCRYPT_R = 8
SCRYPT_P = 1
SCRYPT_DKLEN = 32


class DocumentKind(StrEnum):
    QUALITY_RECORD = "工程材料／設備品質抽驗紀錄表"
    QUALITY_STANDARD = "材料品質標準抽驗表"
    SUBMISSION_CONTROL = "材料設備送審管制總表"
    INSPECTION_CONTROL = "材料設備檢（試）驗管制總表"


OUTPUT_ORDER = (
    DocumentKind.SUBMISSION_CONTROL,
    DocumentKind.QUALITY_STANDARD,
    DocumentKind.QUALITY_RECORD,
    DocumentKind.INSPECTION_CONTROL,
)


@dataclass(frozen=True, slots=True)
class AppSettings:
    password_hash: str
    allow_unsecured_local: bool


@dataclass(frozen=True, slots=True)
class ValidationIssue:
    code: str
    message: str
    document: str = ""
    table: int | None = None
    row: int | None = None
    column: str | None = None
    expected: str | None = None
    actual: str | None = None

    def display(self) -> str:
        location = [self.document] if self.document else []
        if self.table is not None:
            location.append(f"表格 {self.table}")
        if self.row is not None:
            location.append(f"第 {self.row} 列")
        if self.column:
            location.append(self.column)
        detail = f"{' / '.join(location) or '檔案'}：{self.message}"
        if self.expected is not None:
            detail += f"；預期 {self.expected!r}"
        if self.actual is not None:
            detail += f"；實際 {self.actual!r}"
        return detail


@dataclass(frozen=True, slots=True)
class MaterialRecord:
    name: str
    source_row: int
    item_count: int = 1


@dataclass(frozen=True, slots=True)
class MaterialSelection:
    index: int
    name: str
    contract_item: str
    contract_quantity: str
    planned_submission_date: str = ""


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    kind: DocumentKind
    title: str
    materials: tuple[MaterialRecord, ...]


@dataclass(frozen=True, slots=True)
class DocumentValidationResult:
    parsed: ParsedDocument | None
    issues: tuple[ValidationIssue, ...]

    @property
    def valid(self) -> bool:
        return self.parsed is not None and not self.issues


@dataclass(frozen=True, slots=True)
class UploadedSource:
    upload_name: str
    content: bytes


@dataclass(frozen=True, slots=True)
class BatchValidationResult:
    document_results: tuple[DocumentValidationResult, ...]
    issues: tuple[ValidationIssue, ...]


@dataclass(frozen=True, slots=True)
class OpenXmlPackage:
    entries: dict[str, bytes]
    document: etree._Element


class FourFormGenerationError(ValueError):
    pass


# =============================================================================
# 密碼與 Secrets
# =============================================================================

def _derive_password(password: str, salt: bytes, *, n: int, r: int, p: int) -> bytes:
    return hashlib.scrypt(
        password.encode("utf-8"),
        salt=salt,
        n=n,
        r=r,
        p=p,
        dklen=SCRYPT_DKLEN,
    )


def hash_password(password: str) -> str:
    if not password:
        raise ValueError("密碼不可為空")
    salt = secrets.token_bytes(16)
    derived = _derive_password(password, salt, n=SCRYPT_N, r=SCRYPT_R, p=SCRYPT_P)
    salt_text = base64.urlsafe_b64encode(salt).decode("ascii").rstrip("=")
    hash_text = base64.urlsafe_b64encode(derived).decode("ascii").rstrip("=")
    return f"$scrypt$n={SCRYPT_N},r={SCRYPT_R},p={SCRYPT_P}${salt_text}${hash_text}"


def _decode_base64(value: str) -> bytes:
    return base64.urlsafe_b64decode(value + "=" * (-len(value) % 4))


def verify_password(password: str, encoded_hash: str) -> bool:
    if not password or not encoded_hash:
        return False
    try:
        marker, algorithm, params, salt_text, expected_text = encoded_hash.split("$")
        if marker or algorithm != "scrypt":
            return False
        parsed = dict(item.split("=", 1) for item in params.split(","))
        n, r, p = int(parsed["n"]), int(parsed["r"]), int(parsed["p"])
        if (n, r, p) != (SCRYPT_N, SCRYPT_R, SCRYPT_P):
            return False
        actual = _derive_password(password, _decode_base64(salt_text), n=n, r=r, p=p)
        expected = _decode_base64(expected_text)
    except (ValueError, KeyError):
        return False
    return hmac.compare_digest(actual, expected)


def _secret_section(values: Mapping[str, Any], name: str) -> Mapping[str, Any]:
    value = values.get(name, {})
    return value if isinstance(value, Mapping) else {}


def load_settings(values: Mapping[str, Any]) -> AppSettings:
    app = _secret_section(values, "app")
    return AppSettings(
        password_hash=str(app.get("password_hash", "")),
        allow_unsecured_local=app.get("allow_unsecured_local") is True,
    )


def secret_mapping() -> dict[str, Any]:
    try:
        return dict(st.secrets)
    except Exception:
        return {}


# =============================================================================
# 四份正式 Word 的格式契約
# =============================================================================

@dataclass(frozen=True, slots=True)
class TableContract:
    row_count: int | None
    grid_columns: int
    structure_signature: str | None
    header_shapes: tuple[str, ...] = ()
    header_texts: tuple[tuple[str | None, ...], ...] = ()


@dataclass(frozen=True, slots=True)
class DocumentContract:
    kind: DocumentKind
    exact_title: str
    tables: tuple[TableContract, ...]
    page_size: tuple[str, str, str]
    page_margins: tuple[str, str, str, str, str, str, str]


QUALITY_RECORD_CONTRACT = DocumentContract(
    DocumentKind.QUALITY_RECORD,
    "表3-4-工程材料/設備品質抽驗紀錄表",
    (
        TableContract(
            None,
            5,
            None,
            ("1:-|4:-", "1:-|1:-|1:-|2:-", "1:-|1:-|1:-|1:-|1:-"),
            (
                ("工程名稱", None),
                ("材料/設備名稱", None, "檢查日期", None),
                ("抽驗項目", "抽驗標準", "抽驗數量", "抽驗值", "抽驗結果"),
            ),
        ),
        TableContract(
            1,
            2,
            "ff5f1f1f7258acd05cebe0796af718cf594aacb4b85b0201a6cf208494fe5768",
            ("1:-|1:-",),
            (("監造單位主管", "監造現場人員"),),
        ),
    ),
    ("11906", "16838", "portrait"),
    ("1000", "1080", "1000", "1080", "708", "708", "0"),
)

QUALITY_STANDARD_CONTRACT = DocumentContract(
    DocumentKind.QUALITY_STANDARD,
    "材料品質標準抽驗表",
    (
        TableContract(
            None,
            4,
            None,
            ("2:-|1:-|1:-",),
            (("工程項目", "抽查頻率", "管理標準"),),
        ),
    ),
    ("12240", "15840", "portrait"),
    ("1080", "1080", "1080", "1080", "708", "708", "0"),
)

SUBMISSION_CONTROL_CONTRACT = DocumentContract(
    DocumentKind.SUBMISSION_CONTROL,
    "表3-2 材料設備送審管制總表",
    (
        TableContract(
            None,
            15,
            None,
            (
                "10:-|5:-",
                "1:restart|1:restart|1:restart|1:restart|1:restart|1:restart|1:restart|6:-|1:restart|1:restart",
                "1:continue|1:continue|1:continue|1:continue|1:continue|1:continue|1:continue|1:restart|1:restart|2:restart|1:restart|1:restart|1:continue|1:continue",
                "1:continue|1:-|1:continue|1:continue|1:-|1:-|1:continue|1:continue|1:continue|2:continue|1:continue|1:continue|1:-|1:continue",
            ),
            (
                (None, None),
                (
                    "項次", "契約詳細表項次", "契約數量", "是否取樣試驗", "預定\n送審日期",
                    "是否驗廠", "預定\n試驗單位", "送審資料（ˇ）", "審查日期", "備註歸檔編號",
                ),
                ("", "", "", "", "", "", "", "協力廠商資料", "型錄", "相關試驗報告", "樣品", "其他", "", ""),
                ("", "材料(設備)名稱", "", "", "實際\n送審日期", "驗廠日期", "", "", "", "", "", "", "審查結果", ""),
            ),
        ),
    ),
    ("16838", "11906", "landscape"),
    ("1800", "1440", "1800", "1440", "851", "992", "0"),
)

INSPECTION_CONTROL_CONTRACT = DocumentContract(
    DocumentKind.INSPECTION_CONTROL,
    "表3-5材料設備檢(試)驗管制總表",
    (
        TableContract(
            None,
            11,
            None,
            (
                "8:-|3:restart",
                "8:-|3:continue",
                "1:restart|1:-|1:-|1:restart|1:-|1:restart|1:-|2:restart|1:restart|1:restart",
                "1:continue|1:-|1:-|1:continue|1:-|1:continue|1:-|2:continue|1:continue|1:continue",
            ),
            (
                (None, None),
                (None, ""),
                (
                    "項次", "契約詳細表項次", "預定進場日期", "進場數量", "抽樣日期",
                    "規定抽(取)樣頻率", "累積進場數量", "檢(試)驗結果",
                    "檢(試)驗及會同人員", "備註(歸檔編號)",
                ),
                ("", "材料(設備)名稱", "實際進場日期", "", "抽樣數量", "", "累積抽樣數量", "", "", ""),
            ),
        ),
    ),
    ("15840", "12240", "landscape"),
    ("720", "720", "720", "720", "708", "708", "0"),
)

CONTRACTS = {
    contract.exact_title: contract
    for contract in (
        QUALITY_RECORD_CONTRACT,
        QUALITY_STANDARD_CONTRACT,
        SUBMISSION_CONTROL_CONTRACT,
        INSPECTION_CONTROL_CONTRACT,
    )
}


# =============================================================================
# DOCX 安全讀取、格式驗證與材料擷取
# =============================================================================

def _safe_zip_name(name: str) -> bool:
    path = PurePosixPath(name)
    return not path.is_absolute() and ".." not in path.parts and "\\" not in name


def open_docx(content: bytes, document_name: str) -> tuple[OpenXmlPackage | None, tuple[ValidationIssue, ...]]:
    if not content:
        return None, (ValidationIssue("empty_file", "上傳檔案是空的", document_name),)
    if len(content) > MAX_DOCX_BYTES:
        return None, (ValidationIssue("file_too_large", "DOCX 超過 25 MB", document_name),)
    try:
        with ZipFile(BytesIO(content), "r") as archive:
            infos = archive.infolist()
            names = [info.filename for info in infos]
            if len(infos) > MAX_ZIP_ENTRIES:
                raise ValueError("壓縮項目過多")
            if len(names) != len(set(names)) or any(not _safe_zip_name(name) for name in names):
                raise ValueError("壓縮包路徑不安全或重複")
            if any(info.flag_bits & 0x1 for info in infos):
                raise ValueError("不接受加密 DOCX")
            if sum(info.file_size for info in infos) > MAX_UNCOMPRESSED_BYTES:
                raise ValueError("解壓後內容過大")
            if not {"[Content_Types].xml", "word/document.xml"}.issubset(names):
                raise ValueError("缺少 DOCX 必要內容")
            entries = {info.filename: archive.read(info) for info in infos}
    except (BadZipFile, OSError, ValueError) as exc:
        return None, (ValidationIssue("invalid_docx", f"不是可安全讀取的 DOCX：{exc}", document_name),)

    prohibited = [
        name for name in entries
        if name in PROHIBITED_PART_NAMES or name.startswith(PROHIBITED_PART_PREFIXES)
    ]
    if prohibited:
        return None, (ValidationIssue("prohibited_part", "DOCX 含巨集、ActiveX 或嵌入物件", document_name),)

    parser = etree.XMLParser(resolve_entities=False, load_dtd=False, no_network=True, recover=False)
    for name, content_part in entries.items():
        if not name.endswith(".rels"):
            continue
        try:
            root = etree.fromstring(content_part, parser=parser)
        except etree.XMLSyntaxError as exc:
            return None, (ValidationIssue("invalid_relationship", f"關聯檔無法解析：{exc}", document_name),)
        if root.xpath(".//*[local-name()='Relationship' and @TargetMode='External']"):
            return None, (ValidationIssue("external_link", "DOCX 不接受外部連結", document_name),)
    try:
        document = etree.fromstring(entries["word/document.xml"], parser=parser)
    except etree.XMLSyntaxError as exc:
        return None, (ValidationIssue("invalid_xml", f"document.xml 無法解析：{exc}", document_name),)
    return OpenXmlPackage(entries, document), ()


def visible_text(element: etree._Element) -> str:
    parts: list[str] = []
    for node in element.iter():
        if node.tag == W + "t":
            parts.append(node.text or "")
        elif node.tag == W + "tab":
            parts.append("\t")
        elif node.tag in {W + "br", W + "cr"}:
            parts.append("\n")
    return "".join(parts)


def _cell_shape(cell: etree._Element) -> str:
    span = cell.find("./w:tcPr/w:gridSpan", NS)
    merge = cell.find("./w:tcPr/w:vMerge", NS)
    span_value = span.get(W + "val", "1") if span is not None else "1"
    merge_value = "-" if merge is None else merge.get(W + "val", "continue")
    return f"{span_value}:{merge_value}"


def row_shape(row: etree._Element) -> str:
    return "|".join(_cell_shape(cell) for cell in row.findall("./w:tc", NS))


def structure_digest(table: etree._Element) -> str:
    rows = table.findall("./w:tr", NS)
    grids = table.findall("./w:tblGrid/w:gridCol", NS)
    # Word may rewrite table style XML when a file is opened and saved. For the
    # import contract, the critical structure is the grid count plus each row's
    # cell count / horizontal span / vertical merge pattern.
    structure = "grid:" + str(len(grids)) + "\n" + "\n".join(row_shape(row) for row in rows)
    return hashlib.sha256(structure.encode("utf-8")).hexdigest()


def page_size_tuple(page_size: etree._Element | None) -> tuple[str, str, str]:
    if page_size is None:
        return "", "", ""
    width = page_size.get(W + "w", "")
    height = page_size.get(W + "h", "")
    orientation = page_size.get(W + "orient", "")
    if not orientation and width and height and int(width) <= int(height):
        orientation = "portrait"
    return width, height, orientation


def _validate_contract(
    package: OpenXmlPackage,
    document_name: str,
    contract: DocumentContract,
) -> tuple[list[etree._Element], list[ValidationIssue]]:
    issues: list[ValidationIssue] = []
    tables = package.document.findall("./w:body/w:tbl", NS)
    if len(tables) != len(contract.tables):
        issues.append(ValidationIssue(
            "table_count", "表格數量不符合正式格式", document_name,
            expected=str(len(contract.tables)), actual=str(len(tables)),
        ))
    for table_index, table_contract in enumerate(contract.tables):
        if table_index >= len(tables):
            break
        table = tables[table_index]
        rows = table.findall("./w:tr", NS)
        grids = table.findall("./w:tblGrid/w:gridCol", NS)
        if table_contract.row_count is not None and len(rows) != table_contract.row_count:
            issues.append(ValidationIssue(
                "row_count", "列數不符合正式格式", document_name, table_index + 1,
                expected=str(table_contract.row_count), actual=str(len(rows)),
            ))
        if len(grids) != table_contract.grid_columns:
            issues.append(ValidationIssue(
                "grid_count", "欄數不符合正式格式", document_name, table_index + 1,
                expected=str(table_contract.grid_columns), actual=str(len(grids)),
            ))
        actual_signature = structure_digest(table)
        if table_contract.structure_signature is not None and actual_signature != table_contract.structure_signature:
            issues.append(ValidationIssue(
                "structure", "表格、列或合併儲存格結構已變動", document_name, table_index + 1,
                expected=table_contract.structure_signature, actual=actual_signature,
            ))
        for row_index, expected_shape in enumerate(table_contract.header_shapes):
            if row_index < len(rows) and row_shape(rows[row_index]) != expected_shape:
                issues.append(ValidationIssue(
                    "header_shape", "表頭合併儲存格已變動", document_name,
                    table_index + 1, row_index + 1,
                    expected=expected_shape, actual=row_shape(rows[row_index]),
                ))
        for row_index, expected_texts in enumerate(table_contract.header_texts):
            if row_index >= len(rows):
                continue
            cells = rows[row_index].findall("./w:tc", NS)
            for cell_index, expected in enumerate(expected_texts):
                if expected is None:
                    continue
                actual = visible_text(cells[cell_index]) if cell_index < len(cells) else "（無此欄）"
                if actual != expected:
                    issues.append(ValidationIssue(
                        "header_text", "表頭欄位或順序已變動", document_name,
                        table_index + 1, row_index + 1, f"第 {cell_index + 1} 格",
                        expected, actual,
                    ))
                    break

    sections = package.document.findall(".//w:sectPr", NS)
    if len(sections) != 1:
        issues.append(ValidationIssue(
            "section_count", "分節數不符合正式格式", document_name,
            expected="1", actual=str(len(sections)),
        ))
    else:
        page_size = sections[0].find("./w:pgSz", NS)
        margins = sections[0].find("./w:pgMar", NS)
        actual_page_size = page_size_tuple(page_size)
        actual_margins = tuple(
            margins.get(W + key, "") if margins is not None else ""
            for key in ("top", "right", "bottom", "left", "header", "footer", "gutter")
        )
        if actual_page_size != contract.page_size:
            issues.append(ValidationIssue(
                "page_size", "紙張尺寸或方向已變動", document_name,
                expected=str(contract.page_size), actual=str(actual_page_size),
            ))
        if actual_margins != contract.page_margins:
            issues.append(ValidationIssue(
                "page_margin", "頁面邊界已變動", document_name,
                expected=str(contract.page_margins), actual=str(actual_margins),
            ))
    return tables, issues


def _material_issues(records: list[MaterialRecord], document_name: str) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []
    seen: dict[str, int] = {}
    for record in records:
        if not record.name:
            issues.append(ValidationIssue(
                "blank_material", "材料名稱不可空白", document_name, 1, record.source_row, "材料名稱",
            ))
        elif record.name != record.name.strip():
            issues.append(ValidationIssue(
                "material_space", "材料名稱含前置或尾端空白", document_name,
                1, record.source_row, "材料名稱", actual=record.name,
            ))
        elif record.name in seen:
            issues.append(ValidationIssue(
                "duplicate_material", f"材料名稱與第 {seen[record.name]} 列重複", document_name,
                1, record.source_row, "材料名稱", actual=record.name,
            ))
        else:
            seen[record.name] = record.source_row
    if not records:
        issues.append(ValidationIssue(
            "material_count", "正式文件至少必須包含 1 項材料", document_name,
            expected="至少 1", actual=str(len(records)),
        ))
    return issues


def _parse_quality_record(tables: list[etree._Element], document_name: str) -> tuple[list[MaterialRecord], list[ValidationIssue]]:
    rows = tables[0].findall("./w:tr", NS)
    records: list[MaterialRecord] = []
    issues: list[ValidationIssue] = []
    if len(rows) < 6:
        issues.append(ValidationIssue(
            "quality_record_rows", "表格列數不足，至少需要表頭、材料列與尾端紀錄列", document_name, 1,
            expected="至少 6", actual=str(len(rows)),
        ))
        return records, issues
    for row_index, row in enumerate(rows[-2:], start=len(rows) - 1):
        if row_shape(row) != "5:-":
            issues.append(ValidationIssue(
                "quality_record_footer", "尾端處理方式或特殊紀錄列格式不正確", document_name,
                1, row_index, expected="5:-", actual=row_shape(row),
            ))
    for row_index, row in enumerate(rows[3:-2], start=4):
        cells = row.findall("./w:tc", NS)
        shape = row_shape(row)
        if shape != "1:-|1:-|1:-|1:-|1:-":
            issues.append(ValidationIssue(
                "data_columns", "材料列必須有五格且不可合併", document_name,
                1, row_index, expected="1:-|1:-|1:-|1:-|1:-", actual=shape,
            ))
            continue
        records.append(MaterialRecord(visible_text(cells[0]), row_index))
        if visible_text(cells[4]) != "□合格□不合格":
            issues.append(ValidationIssue(
                "result_placeholder", "抽驗結果格式已變動", document_name,
                1, row_index, "抽驗結果", "□合格□不合格", visible_text(cells[4]),
            ))
    return records, issues


def _flat_cell_text(cell: etree._Element) -> str:
    return visible_text(cell).replace("\n", "").strip()


def _normalize_standard_token(value: str) -> str:
    return value.replace("²", "2").strip()


def _parenthesized_token(value: str) -> str:
    match = re.search(r"[（(]([^）)]+)[）)]", value)
    return _normalize_standard_token(match.group(1)) if match else ""


def _quality_standard_group_name(group_text: str, item_text: str) -> str:
    group = group_text.replace("\n", "").strip()
    token = _parenthesized_token(item_text)
    if group == "預拌混凝土" and token:
        return f"{token}{group}"
    if group == "鋼筋" and token:
        return f"{group}，{token}"
    return group


def _quality_standard_record_name(row: etree._Element) -> str:
    cells = row.findall("./w:tc", NS)
    shape = row_shape(row)
    if shape == "2:-|1:-|1:-":
        return _flat_cell_text(cells[0])
    if shape == "1:restart|1:-|1:-|1:-":
        return _quality_standard_group_name(visible_text(cells[0]), visible_text(cells[1]))
    return ""


def _parse_quality_standard(tables: list[etree._Element], document_name: str) -> tuple[list[MaterialRecord], list[ValidationIssue]]:
    rows = tables[0].findall("./w:tr", NS)
    records: list[MaterialRecord] = []
    issues: list[ValidationIssue] = []
    index = 1
    while index < len(rows):
        row = rows[index]
        shape = row_shape(row)
        cells = row.findall("./w:tc", NS)
        source_row = index + 1
        if shape == "2:-|1:-|1:-":
            records.append(MaterialRecord(_quality_standard_record_name(row), source_row, 1))
            index += 1
        elif shape == "1:restart|1:-|1:-|1:-":
            name = _quality_standard_record_name(row)
            count = 1
            index += 1
            while index < len(rows) and row_shape(rows[index]) == "1:continue|1:-|1:-|1:-":
                count += 1
                index += 1
            records.append(MaterialRecord(name, source_row, count))
        else:
            issues.append(ValidationIssue(
                "quality_standard_shape", "材料列合併結構不正確", document_name,
                1, source_row, expected="正式單列或多列格式", actual=shape,
            ))
            index += 1
    return records, issues


def _parse_submission(tables: list[etree._Element], document_name: str) -> tuple[list[MaterialRecord], list[ValidationIssue]]:
    rows = tables[0].findall("./w:tr", NS)
    records: list[MaterialRecord] = []
    issues: list[ValidationIssue] = []
    top_shape = "1:restart|1:-|1:restart|1:restart|1:-|1:-|1:-|1:restart|1:restart|2:restart|1:restart|1:restart|1:-|1:restart"
    bottom_shape = "1:continue|1:-|1:continue|1:continue|1:-|1:-|1:-|1:continue|1:continue|2:continue|1:continue|1:continue|1:-|1:continue"
    blank_started = False
    if (len(rows) - 4) % 2:
        issues.append(ValidationIssue(
            "submission_row_pairs", "材料資料列必須是每項材料兩列一組", document_name,
            1, expected="表頭後偶數列", actual=str(len(rows) - 4),
        ))
    for offset in range(4, len(rows) - 1, 2):
        top, bottom = rows[offset], rows[offset + 1]
        if row_shape(top) != top_shape or row_shape(bottom) != bottom_shape:
            issues.append(ValidationIssue(
                "submission_pair", "材料雙列合併結構不正確", document_name,
                1, offset + 1,
            ))
            continue
        cells = bottom.findall("./w:tc", NS)
        name = visible_text(cells[1])
        if not name:
            blank_started = True
            continue
        if blank_started:
            issues.append(ValidationIssue(
                "material_after_blank", "空白預留列後不可再出現材料", document_name,
                1, offset + 2, "材料名稱", actual=name,
            ))
        records.append(MaterialRecord(name, offset + 2))
    return records, issues


def _parse_inspection(tables: list[etree._Element], document_name: str) -> tuple[list[MaterialRecord], list[ValidationIssue]]:
    rows = tables[0].findall("./w:tr", NS)
    records: list[MaterialRecord] = []
    issues: list[ValidationIssue] = []
    top_shape = "1:restart|1:-|1:-|1:restart|1:-|1:restart|1:restart|2:restart|1:restart|1:restart"
    bottom_shape = "1:continue|1:-|1:-|1:continue|1:-|1:continue|1:continue|2:continue|1:continue|1:continue"
    if (len(rows) - 4) % 2:
        issues.append(ValidationIssue(
            "inspection_row_pairs", "材料資料列必須是每項材料兩列一組", document_name,
            1, expected="表頭後偶數列", actual=str(len(rows) - 4),
        ))
    for offset in range(4, len(rows) - 1, 2):
        top, bottom = rows[offset], rows[offset + 1]
        if row_shape(top) != top_shape or row_shape(bottom) != bottom_shape:
            issues.append(ValidationIssue(
                "inspection_pair", "材料雙列合併結構不正確", document_name,
                1, offset + 1,
            ))
            continue
        name = visible_text(bottom.findall("./w:tc", NS)[1])
        records.append(MaterialRecord(name, offset + 2))
    return records, issues


def parse_document(content: bytes, document_name: str) -> DocumentValidationResult:
    package, package_issues = open_docx(content, document_name)
    if package is None:
        return DocumentValidationResult(None, package_issues)
    for code, path in {
        "tracked_insert": ".//w:ins",
        "tracked_delete": ".//w:del",
        "alt_chunk": ".//w:altChunk",
        "content_control": ".//w:sdt",
    }.items():
        if package.document.find(path, NS) is not None:
            return DocumentValidationResult(
                None,
                (ValidationIssue(code, "DOCX 含不允許的修訂或動態內容", document_name),),
            )

    titles = [
        visible_text(paragraph)
        for paragraph in package.document.findall("./w:body/w:p", NS)
        if visible_text(paragraph)
    ]
    matched = [title for title in titles if title in CONTRACTS]
    if len(matched) != 1:
        return DocumentValidationResult(
            None,
            (ValidationIssue(
                "unsupported_title", "找不到唯一的正式文件標題", document_name,
                expected="、".join(CONTRACTS), actual=" | ".join(titles) or "（無標題）",
            ),),
        )
    title = matched[0]
    contract = CONTRACTS[title]
    tables, issues = _validate_contract(package, document_name, contract)
    if len(tables) < len(contract.tables):
        return DocumentValidationResult(None, tuple(issues))

    if contract.kind == DocumentKind.QUALITY_RECORD:
        records, record_issues = _parse_quality_record(tables, document_name)
    elif contract.kind == DocumentKind.QUALITY_STANDARD:
        records, record_issues = _parse_quality_standard(tables, document_name)
    elif contract.kind == DocumentKind.SUBMISSION_CONTROL:
        records, record_issues = _parse_submission(tables, document_name)
    else:
        records, record_issues = _parse_inspection(tables, document_name)
    issues.extend(record_issues)
    issues.extend(_material_issues(records, document_name))

    if issues:
        return DocumentValidationResult(None, tuple(issues))
    return DocumentValidationResult(
        ParsedDocument(contract.kind, title, tuple(records)),
        (),
    )


def validate_import_batch(sources: list[UploadedSource]) -> BatchValidationResult:
    issues: list[ValidationIssue] = []
    if len(sources) != 4:
        issues.append(ValidationIssue(
            "upload_count", "必須一次上傳四份 DOCX", expected="4", actual=str(len(sources)),
        ))
    results = tuple(parse_document(source.content, source.upload_name) for source in sources)
    seen: set[DocumentKind] = set()
    parsed_by_kind: dict[DocumentKind, ParsedDocument] = {}
    for source, result in zip(sources, results, strict=True):
        if not result.valid or result.parsed is None:
            issues.extend(result.issues)
            continue
        if result.parsed.kind in seen:
            issues.append(ValidationIssue(
                "duplicate_kind", "四份文件中出現重複表格種類", source.upload_name,
                actual=result.parsed.kind.value,
            ))
        seen.add(result.parsed.kind)
        parsed_by_kind[result.parsed.kind] = result.parsed
    for kind in DocumentKind:
        if kind not in seen:
            issues.append(ValidationIssue("missing_kind", "缺少此正式文件", kind.value))
    if set(parsed_by_kind) == set(DocumentKind):
        reference = parsed_by_kind[DocumentKind.SUBMISSION_CONTROL]
        reference_count = len(reference.materials)
        for kind, parsed in parsed_by_kind.items():
            if kind == DocumentKind.SUBMISSION_CONTROL:
                continue
            if len(parsed.materials) != reference_count:
                issues.append(ValidationIssue(
                    "material_count_mismatch", "材料數量與材料設備送審管制總表不一致",
                    parsed.title,
                    expected=f"{reference_count} 項",
                    actual=f"{len(parsed.materials)} 項",
                ))
    return BatchValidationResult(results, tuple(issues))


# =============================================================================
# 依選材順序篩選四表並合併成一份 Word
# =============================================================================

def _set_paragraph_alignment(paragraph_properties: etree._Element, alignment: str) -> None:
    for existing in paragraph_properties.findall("./w:jc", NS):
        paragraph_properties.remove(existing)
    justify = etree.SubElement(paragraph_properties, W + "jc")
    justify.set(W + "val", alignment)


def _set_cell_text(cell: etree._Element, value: str, *, alignment: str | None = None) -> None:
    paragraphs = cell.findall("./w:p", NS)
    template = paragraphs[0] if paragraphs else None
    paragraph_properties = template.find("./w:pPr", NS) if template is not None else None
    first_run = template.find(".//w:r", NS) if template is not None else None
    run_properties = first_run.find("./w:rPr", NS) if first_run is not None else None
    for paragraph in paragraphs:
        cell.remove(paragraph)
    paragraph = etree.SubElement(cell, W + "p")
    if paragraph_properties is not None:
        paragraph_properties = deepcopy(paragraph_properties)
        if alignment:
            _set_paragraph_alignment(paragraph_properties, alignment)
        paragraph.append(paragraph_properties)
    elif alignment:
        paragraph_properties = etree.SubElement(paragraph, W + "pPr")
        _set_paragraph_alignment(paragraph_properties, alignment)
    run = etree.SubElement(paragraph, W + "r")
    if run_properties is not None:
        run.append(deepcopy(run_properties))
    for line_index, line in enumerate(value.split("\n")):
        if line_index:
            etree.SubElement(run, W + "br")
        text = etree.SubElement(run, W + "t")
        if line.startswith(" ") or line.endswith(" "):
            text.set(XML_SPACE, "preserve")
        text.text = line


def _source_default_run_format(entries: dict[str, bytes]) -> tuple[dict[str, str], str | None]:
    styles_content = entries.get("word/styles.xml")
    if not styles_content:
        return {}, None
    styles = etree.fromstring(styles_content)
    run_properties = styles.find("./w:docDefaults/w:rPrDefault/w:rPr", NS)
    if run_properties is None:
        return {}, None
    fonts = run_properties.find("./w:rFonts", NS)
    font_values: dict[str, str] = {}
    if fonts is not None:
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            value = fonts.get(W + key)
            if value:
                font_values[key] = value
    size = run_properties.find("./w:sz", NS)
    return font_values, size.get(W + "val") if size is not None else None


def _ensure_run_properties(run: etree._Element) -> etree._Element:
    run_properties = run.find("./w:rPr", NS)
    if run_properties is None:
        run_properties = etree.Element(W + "rPr")
        run.insert(0, run_properties)
    return run_properties


def _materialize_source_run_defaults(entries: dict[str, bytes], document: etree._Element) -> None:
    default_fonts, default_size = _source_default_run_format(entries)
    # Some official forms rely on Word's implicit 11pt default instead of writing
    # w:sz into styles.xml. Materializing it prevents later merged sections from
    # inheriting the first document's default size.
    if default_fonts and default_size is None:
        default_size = "22"
    if not default_fonts and not default_size:
        return
    for run in document.findall(".//w:r", NS):
        if run.find("./w:t", NS) is None:
            continue
        run_properties = _ensure_run_properties(run)
        if default_fonts:
            fonts = run_properties.find("./w:rFonts", NS)
            if fonts is None:
                fonts = etree.Element(W + "rFonts")
                run_properties.insert(0, fonts)
            for key, value in default_fonts.items():
                if fonts.get(W + key) is None:
                    fonts.set(W + key, value)
        if default_size:
            if run_properties.find("./w:sz", NS) is None:
                size = etree.Element(W + "sz")
                size.set(W + "val", default_size)
                run_properties.append(size)
            if run_properties.find("./w:szCs", NS) is None:
                size_cs = etree.Element(W + "szCs")
                size_cs.set(W + "val", default_size)
                run_properties.append(size_cs)


def _serialize_package(
    entries: dict[str, bytes],
    document: etree._Element,
    *,
    materialize_run_defaults: bool = False,
) -> bytes:
    if materialize_run_defaults:
        _materialize_source_run_defaults(entries, document)
    updated = dict(entries)
    updated["word/document.xml"] = etree.tostring(
        document, xml_declaration=True, encoding="UTF-8", standalone=True,
    )
    output = BytesIO()
    with ZipFile(output, "w", compression=ZIP_DEFLATED) as archive:
        for name, content in updated.items():
            archive.writestr(name, content)
    return output.getvalue()


def _open_generation_source(
    kind: DocumentKind,
    source: UploadedSource,
) -> tuple[ParsedDocument, dict[str, bytes], etree._Element]:
    result = parse_document(source.content, source.upload_name)
    if not result.valid or result.parsed is None:
        details = "\n".join(issue.display() for issue in result.issues)
        raise FourFormGenerationError(f"{source.upload_name} 無法通過格式驗證：\n{details}")
    if result.parsed.kind != kind:
        raise FourFormGenerationError(f"{source.upload_name} 的表格種類不正確。")
    package, issues = open_docx(source.content, source.upload_name)
    if package is None:
        raise FourFormGenerationError("\n".join(issue.display() for issue in issues))
    return result.parsed, package.entries, deepcopy(package.document)


def _replace_rows(table: etree._Element, rows: list[etree._Element]) -> None:
    for row in table.findall("./w:tr", NS):
        table.remove(row)
    for row in rows:
        table.append(deepcopy(row))


def _filter_quality_record(
    entries: dict[str, bytes],
    document: etree._Element,
    indices: tuple[int, ...],
    project_name: str,
) -> bytes:
    tables = document.findall("./w:body/w:tbl", NS)
    rows = tables[0].findall("./w:tr", NS)
    header = [deepcopy(row) for row in rows[:3]]
    _set_cell_text(header[0].findall("./w:tc", NS)[1], project_name)
    selected = []
    for index in indices:
        row = deepcopy(rows[3 + index])
        cells = row.findall("./w:tc", NS)
        _set_cell_text(cells[3], "")
        _set_cell_text(cells[4], "□合格□不合格")
        selected.append(row)
    _replace_rows(tables[0], header + selected + [deepcopy(row) for row in rows[-2:]])
    return _serialize_package(entries, document, materialize_run_defaults=True)


def _filter_quality_standard(
    parsed: ParsedDocument,
    entries: dict[str, bytes],
    document: etree._Element,
    indices: tuple[int, ...],
) -> bytes:
    table = document.find("./w:body/w:tbl", NS)
    if table is None:
        raise FourFormGenerationError("材料品質標準抽驗表找不到表格。")
    rows = table.findall("./w:tr", NS)
    output_rows = [deepcopy(rows[0])]
    for index in indices:
        record = parsed.materials[index]
        start = record.source_row - 1
        output_rows.extend(deepcopy(row) for row in rows[start : start + record.item_count])
    _replace_rows(table, output_rows)
    return _serialize_package(entries, document, materialize_run_defaults=True)


def _filter_submission(
    entries: dict[str, bytes],
    document: etree._Element,
    selections: tuple[MaterialSelection, ...],
    project_name: str,
) -> bytes:
    table = document.find("./w:body/w:tbl", NS)
    if table is None:
        raise FourFormGenerationError("材料設備送審管制總表找不到表格。")
    rows = table.findall("./w:tr", NS)
    header = [deepcopy(row) for row in rows[:4]]
    header_cells = header[0].findall("./w:tc", NS)
    _set_cell_text(header_cells[0], f"工程名稱：{project_name}\n主辦單位：")
    _set_cell_text(header_cells[1], "表單編號：")
    output_rows = header
    for sequence, selection in enumerate(selections, start=1):
        top = deepcopy(rows[4 + 2 * selection.index])
        bottom = deepcopy(rows[5 + 2 * selection.index])
        top_cells = top.findall("./w:tc", NS)
        bottom_cells = bottom.findall("./w:tc", NS)
        _set_cell_text(top_cells[0], f"{sequence}.")
        _set_cell_text(top_cells[1], selection.contract_item)
        _set_cell_text(top_cells[2], selection.contract_quantity)
        _set_cell_text(top_cells[4], selection.planned_submission_date)
        for cell_index in (12, 13):
            _set_cell_text(top_cells[cell_index], "")
        for cell_index in (4, 5, 12, 13):
            _set_cell_text(bottom_cells[cell_index], "")
        output_rows.extend((top, bottom))
    _replace_rows(table, output_rows)
    return _serialize_package(entries, document, materialize_run_defaults=True)


def _filter_inspection(
    entries: dict[str, bytes],
    document: etree._Element,
    selections: tuple[MaterialSelection, ...],
    project_name: str,
) -> bytes:
    table = document.find("./w:body/w:tbl", NS)
    if table is None:
        raise FourFormGenerationError("材料設備檢（試）驗管制總表找不到表格。")
    rows = table.findall("./w:tr", NS)
    header = [deepcopy(row) for row in rows[:4]]
    first = header[0].findall("./w:tc", NS)
    second = header[1].findall("./w:tc", NS)
    _set_cell_text(first[0], f"▪工程名稱：{project_name}")
    _set_cell_text(first[1], "表單編號：")
    _set_cell_text(second[0], "▪主辦機關：")
    _set_cell_text(second[1], "")
    output_rows = header
    for sequence, selection in enumerate(selections, start=1):
        top = deepcopy(rows[4 + 2 * selection.index])
        bottom = deepcopy(rows[5 + 2 * selection.index])
        top_cells = top.findall("./w:tc", NS)
        bottom_cells = bottom.findall("./w:tc", NS)
        _set_cell_text(top_cells[0], f"{sequence}.")
        _set_cell_text(top_cells[1], selection.contract_item, alignment="center")
        _set_cell_text(top_cells[2], selection.planned_submission_date)
        _set_cell_text(top_cells[3], selection.contract_quantity, alignment="center")
        for cell_index in (4, 6, 7, 8, 9):
            _set_cell_text(top_cells[cell_index], "")
        for cell_index in (0, 2, 3, 4, 5, 6, 7, 8, 9):
            _set_cell_text(bottom_cells[cell_index], "")
        output_rows.extend((top, bottom))
    _replace_rows(table, output_rows)
    return _serialize_package(entries, document, materialize_run_defaults=True)


def _add_section_break(content: bytes) -> bytes:
    package, issues = open_docx(content, "合併前文件")
    if package is None:
        raise FourFormGenerationError("\n".join(issue.display() for issue in issues))
    document = deepcopy(package.document)
    body = document.find("./w:body", NS)
    section = body.find("./w:sectPr", NS) if body is not None else None
    if body is None or section is None:
        raise FourFormGenerationError("合併前文件缺少分節設定。")
    section_copy = deepcopy(section)
    section_type = section_copy.find("./w:type", NS)
    if section_type is None:
        section_type = etree.Element(W + "type")
        section_copy.insert(0, section_type)
    section_type.set(W + "val", "nextPage")
    paragraph = etree.Element(W + "p")
    paragraph_properties = etree.SubElement(paragraph, W + "pPr")
    paragraph_properties.append(section_copy)
    body.insert(body.index(section), paragraph)
    return _serialize_package(package.entries, document)


def _merge_parts(parts: list[bytes]) -> bytes:
    prepared = [_add_section_break(part) for part in parts[:-1]] + [parts[-1]]
    composer = Composer(Document(BytesIO(prepared[0])))
    for part in prepared[1:]:
        composer.append(Document(BytesIO(part)))
    output = BytesIO()
    composer.save(output)

    merged, merged_issues = open_docx(output.getvalue(), "產出文件")
    final, final_issues = open_docx(parts[-1], "最後一表")
    if merged is None or final is None:
        raise FourFormGenerationError("\n".join(issue.display() for issue in merged_issues + final_issues))
    document = deepcopy(merged.document)
    body = document.find("./w:body", NS)
    old_final = body.find("./w:sectPr", NS) if body is not None else None
    new_final = final.document.find("./w:body/w:sectPr", NS)
    if body is None or old_final is None or new_final is None:
        raise FourFormGenerationError("產出文件缺少最終分節設定。")
    body.replace(old_final, deepcopy(new_final))
    return _serialize_package(merged.entries, document)


def _verify_output(
    content: bytes,
    selections: tuple[MaterialSelection, ...],
    parsed: dict[DocumentKind, ParsedDocument],
) -> None:
    package, issues = open_docx(content, "產出文件")
    if package is None:
        raise FourFormGenerationError("\n".join(issue.display() for issue in issues))
    body = package.document.find("./w:body", NS)
    if body is None:
        raise FourFormGenerationError("產出文件缺少本文。")
    expected_titles = [CONTRACTS_BY_KIND[kind].exact_title for kind in OUTPUT_ORDER]
    titles = [
        visible_text(paragraph)
        for paragraph in body.findall("./w:p", NS)
        if visible_text(paragraph) in expected_titles
    ]
    if titles != expected_titles:
        raise FourFormGenerationError("四表標題或順序不正確，已停止輸出。")

    tables = body.findall("./w:tbl", NS)
    indices = tuple(selection.index for selection in selections)
    selected_count = len(selections)
    standard_rows = 1 + sum(parsed[DocumentKind.QUALITY_STANDARD].materials[index].item_count for index in indices)
    expected_rows = [
        4 + 2 * selected_count,
        standard_rows,
        5 + selected_count,
        1,
        4 + 2 * selected_count,
    ]
    actual_rows = [len(table.findall("./w:tr", NS)) for table in tables]
    if len(tables) != 5 or actual_rows != expected_rows:
        raise FourFormGenerationError("四表資料列數驗證失敗，已停止輸出。")

    expected_names = {
        kind: [parsed[kind].materials[index].name for index in indices]
        for kind in OUTPUT_ORDER
    }
    submission_rows = tables[0].findall("./w:tr", NS)
    standard_table_rows = tables[1].findall("./w:tr", NS)
    quality_rows = tables[2].findall("./w:tr", NS)
    inspection_rows = tables[4].findall("./w:tr", NS)
    standard_names: list[str] = []
    standard_row = 1
    for index in indices:
        standard_names.append(_quality_standard_record_name(standard_table_rows[standard_row]))
        standard_row += parsed[DocumentKind.QUALITY_STANDARD].materials[index].item_count
    actual_names = {
        DocumentKind.QUALITY_RECORD: [
            visible_text(quality_rows[3 + offset].findall("./w:tc", NS)[0])
            for offset in range(selected_count)
        ],
        DocumentKind.QUALITY_STANDARD: standard_names,
        DocumentKind.SUBMISSION_CONTROL: [
            visible_text(submission_rows[5 + 2 * offset].findall("./w:tc", NS)[1])
            for offset in range(selected_count)
        ],
        DocumentKind.INSPECTION_CONTROL: [
            visible_text(inspection_rows[5 + 2 * offset].findall("./w:tc", NS)[1])
            for offset in range(selected_count)
        ],
    }
    if actual_names != expected_names:
        raise FourFormGenerationError("四表材料內容或順序驗證失敗，已停止輸出。")

    for offset, selection in enumerate(selections):
        submission_cells = submission_rows[4 + 2 * offset].findall("./w:tc", NS)
        inspection_cells = inspection_rows[4 + 2 * offset].findall("./w:tc", NS)
        actual_submission_values = tuple(
            visible_text(submission_cells[index]) for index in (1, 2, 4)
        )
        expected_submission_values = (
            selection.contract_item,
            selection.contract_quantity,
            selection.planned_submission_date,
        )
        if actual_submission_values != expected_submission_values:
            raise FourFormGenerationError("送審管制總表的契約資料寫入失敗，已停止輸出。")
        actual_inspection_values = (
            visible_text(inspection_cells[1]),
            visible_text(inspection_cells[2]),
            visible_text(inspection_cells[3]),
        )
        if actual_inspection_values != (
            selection.contract_item,
            selection.planned_submission_date,
            selection.contract_quantity,
        ):
            raise FourFormGenerationError("檢（試）驗管制總表的契約項次、預定進場日期或進場數量寫入失敗，已停止輸出。")

    sections = package.document.findall(".//w:sectPr", NS)
    expected_sizes = [CONTRACTS_BY_KIND[kind].page_size for kind in OUTPUT_ORDER]
    actual_sizes = []
    for section in sections:
        actual_sizes.append(page_size_tuple(section.find("./w:pgSz", NS)))
    if len(sections) != 4 or actual_sizes != expected_sizes:
        raise FourFormGenerationError("四表分節、紙張尺寸或方向驗證失敗，已停止輸出。")


CONTRACTS_BY_KIND = {contract.kind: contract for contract in CONTRACTS.values()}


def generate_four_form_docx(
    sources: dict[DocumentKind, UploadedSource],
    project_name: str,
    selected_materials: list[MaterialSelection] | tuple[MaterialSelection, ...],
) -> bytes:
    project_name = project_name.strip()
    selections = tuple(selected_materials)
    indices = tuple(selection.index for selection in selections)
    if not project_name:
        raise FourFormGenerationError("工程名稱不可空白。")
    if set(sources) != set(DocumentKind):
        raise FourFormGenerationError("必須提供四份已驗證的正式 Word。")
    if not selections:
        raise FourFormGenerationError("至少選擇一項材料。")
    if len(indices) != len(set(indices)):
        raise FourFormGenerationError("同一材料不可重複選取。")

    opened = {kind: _open_generation_source(kind, sources[kind]) for kind in OUTPUT_ORDER}
    parsed = {kind: opened[kind][0] for kind in OUTPUT_ORDER}
    submission_materials = parsed[DocumentKind.SUBMISSION_CONTROL].materials
    if any(index < 0 or index >= len(submission_materials) for index in indices):
        raise FourFormGenerationError("材料超出材料設備送審管制總表的範圍。")
    for kind, parsed_document in parsed.items():
        if len(parsed_document.materials) != len(submission_materials):
            raise FourFormGenerationError(f"{kind.value} 的材料數量與材料設備送審管制總表不一致。")
    for selection in selections:
        if submission_materials[selection.index].name != selection.name:
            raise FourFormGenerationError("材料名稱與材料設備送審管制總表不一致。")
        if not selection.contract_item.strip() or not selection.contract_quantity.strip():
            raise FourFormGenerationError("每項材料都必須填寫契約詳細表項次與契約數量。")
    parts = [
        _filter_submission(opened[DocumentKind.SUBMISSION_CONTROL][1], opened[DocumentKind.SUBMISSION_CONTROL][2], selections, project_name),
        _filter_quality_standard(parsed[DocumentKind.QUALITY_STANDARD], opened[DocumentKind.QUALITY_STANDARD][1], opened[DocumentKind.QUALITY_STANDARD][2], indices),
        _filter_quality_record(opened[DocumentKind.QUALITY_RECORD][1], opened[DocumentKind.QUALITY_RECORD][2], indices, project_name),
        _filter_inspection(opened[DocumentKind.INSPECTION_CONTROL][1], opened[DocumentKind.INSPECTION_CONTROL][2], selections, project_name),
    ]
    output = _merge_parts(parts)
    _verify_output(output, selections, parsed)
    return output


# =============================================================================
# Streamlit 網頁介面
# =============================================================================

def apply_theme() -> None:
    st.markdown(
        """
        <style>
        .stApp {
            background-color: #fffaf6;
            background-image:
                radial-gradient(circle at 28px 30px, rgba(184, 210, 224, .30) 0 18px, transparent 19px),
                radial-gradient(circle at 92px 62px, rgba(248, 190, 199, .28) 0 22px, transparent 23px),
                radial-gradient(circle at 148px 28px, rgba(245, 218, 151, .30) 0 15px, transparent 16px);
            background-size: 180px 110px;
        }
        [data-testid="stHeader"], .stAppHeader {
            background: transparent !important;
            height: 0 !important;
            min-height: 0 !important;
        }
        [data-testid="stHeader"] [data-testid="stToolbar"] { display: none !important; }
        [data-testid="stExpandSidebarButton"] {
            position: fixed !important;
            top: .65rem !important;
            left: .65rem !important;
            z-index: 1000000 !important;
            background: rgba(255, 251, 247, .96) !important;
            border: 1px solid rgba(94, 76, 67, .18) !important;
            border-radius: 8px !important;
        }
        [data-testid="stSidebarCollapseButton"] { display: none !important; }
        [data-testid="stSidebar"] {
            transform: translateX(0) !important;
            visibility: visible !important;
            width: 310px !important;
            min-width: 310px;
            max-width: 310px;
            background: rgba(255, 251, 247, .94);
            border-right: 1px solid rgba(94, 76, 67, .16);
        }
        [data-testid="stSidebar"] [data-testid="stVerticalBlock"] { gap: .75rem; }
        [data-testid="stMainBlockContainer"] {
            max-width: 1050px;
            padding-top: 2rem;
            padding-bottom: 2rem;
        }
        h1, h2, h3, p, label { color: #222936; }
        [data-baseweb="input"] > div,
        [data-baseweb="select"] > div,
        [data-testid="stFileUploaderDropzone"] {
            background: rgba(243, 246, 251, .94);
            border-radius: 10px;
        }
        [data-testid="stAlert"] { background: rgba(231, 240, 253, .90); border-radius: 10px; }
        .stButton > button[kind="primary"] {
            background: #ff4b50;
            border: 0;
            border-radius: 10px;
            min-height: 2.7rem;
            font-weight: 700;
        }
        .sidebar-heading {
            color: #222936;
            font-size: 1.55rem;
            font-weight: 800;
            margin-bottom: -.35rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def require_login(settings: AppSettings) -> bool:
    if st.session_state.get("authenticated"):
        with st.sidebar:
            if st.button("登出", use_container_width=True):
                st.session_state.clear()
                st.rerun()
        return True
    has_password = bool(settings.password_hash and "REPLACE_WITH" not in settings.password_hash)
    if not has_password and settings.allow_unsecured_local:
        return True
    st.title(APP_NAME)
    if not has_password:
        st.error("尚未設定登入密碼。")
        st.markdown("本機請修改 `.streamlit/secrets.toml`；Streamlit Cloud 請修改 App Settings → Secrets。")
        st.code(r".\.venv\Scripts\python.exe streamlit_app.py --hash-password", language="powershell")
        return False
    failures = int(st.session_state.get("login_failures", 0))
    remaining = max(0, int(float(st.session_state.get("login_blocked_until", 0.0)) - time.monotonic()))
    if remaining:
        st.error(f"登入失敗次數過多，請在 {remaining} 秒後重試。")
        return False
    with st.form("login_form", clear_on_submit=True):
        password = st.text_input("共用密碼", type="password", autocomplete="current-password")
        submitted = st.form_submit_button("登入", use_container_width=True, type="primary")
    if submitted:
        if verify_password(password, settings.password_hash):
            st.session_state["authenticated"] = True
            st.session_state["login_failures"] = 0
            st.rerun()
        failures += 1
        st.session_state["login_failures"] = failures
        if failures >= MAX_LOGIN_FAILURES:
            st.session_state["login_blocked_until"] = time.monotonic() + LOGIN_COOLDOWN_SECONDS
            st.session_state["login_failures"] = 0
        st.error("密碼錯誤。")
    return False


def render_word_import(settings: AppSettings) -> tuple[str, ...]:
    with st.sidebar:
        st.markdown('<div class="sidebar-heading">Word 檔案匯入</div>', unsafe_allow_html=True)
        st.caption("一次上傳 4 個 Word 檔案")
        uploads = st.file_uploader(
            "選擇四份正式 DOCX",
            type=["docx"],
            accept_multiple_files=True,
            label_visibility="collapsed",
            help="系統讀取文件內的正式表格，不使用檔名判斷。",
        )
        upload_count = len(uploads) if uploads else 0
        st.caption(f"已選擇 {upload_count} / 4 份")
        import_clicked = st.button(
            "匯入 Word 檔案",
            type="primary",
            use_container_width=True,
            disabled=upload_count != 4,
        )
        if import_clicked:
            for key in (
                "imported_materials",
                "imported_sources",
                "generated_docx",
                "generated_fingerprint",
                "planned_submission_month",
            ):
                st.session_state.pop(key, None)
            for key in list(st.session_state):
                if key.startswith((
                    "material_slot_",
                    "contract_item_",
                    "contract_quantity_",
                )):
                    del st.session_state[key]
            sources = [UploadedSource(upload.name, upload.getvalue()) for upload in uploads]
            validation = validate_import_batch(sources)
            parsed_by_kind: dict[DocumentKind, ParsedDocument] = {}
            source_by_kind: dict[DocumentKind, UploadedSource] = {}
            for source, result in zip(sources, validation.document_results, strict=True):
                if result.valid and result.parsed is not None and result.parsed.kind not in parsed_by_kind:
                    parsed_by_kind[result.parsed.kind] = result.parsed
                    source_by_kind[result.parsed.kind] = source
            if set(parsed_by_kind) != set(DocumentKind):
                st.error("有 Word 檔案無法讀取，請確認四份檔案格式。")
                st.download_button(
                    "下載讀取錯誤報告",
                    data="\n".join(issue.display() for issue in validation.issues).encode("utf-8-sig"),
                    file_name="Word讀取錯誤報告.txt",
                    mime="text/plain",
                    use_container_width=True,
                )
            else:
                submission = parsed_by_kind[DocumentKind.SUBMISSION_CONTROL]
                st.session_state["imported_materials"] = tuple(item.name for item in submission.materials)
                st.session_state["imported_sources"] = {
                    kind.name: (source.upload_name, source.content)
                    for kind, source in source_by_kind.items()
                }
                st.session_state["material_slot_count"] = 8
                st.success(f"四份 Word 已讀取，共 {len(submission.materials)} 項材料。")
                st.caption("材料選單依材料設備送審管制總表的原始順序顯示，本次資料不會永久保存。")
        if settings.allow_unsecured_local and not settings.password_hash:
            st.caption("本機開發模式")
    return tuple(st.session_state.get("imported_materials", ()))


def selected_material_names(materials: tuple[str, ...], slot_count: int) -> list[str]:
    return [
        value
        for index in range(slot_count)
        if (value := st.session_state.get(f"material_slot_{index}")) in materials
    ]


def material_sort_key(name: str) -> tuple[tuple[int, int | str], ...]:
    parts = re.split(r"(\d+)", name.casefold())
    return tuple(
        (0, int(part)) if part.isdigit() else (1, part)
        for part in parts
        if part
    )


def add_material_slot() -> None:
    current = max(8, int(st.session_state.get("material_slot_count", 8)))
    material_count = len(st.session_state.get("imported_materials", ()))
    st.session_state["material_slot_count"] = min(current + 1, material_count)


def material_selections(
    materials: tuple[str, ...],
    slot_count: int,
    planned_submission_month: str,
) -> list[MaterialSelection]:
    name_to_index = {name: index for index, name in enumerate(materials)}
    selections: list[MaterialSelection] = []
    for slot_index in range(slot_count):
        name = st.session_state.get(f"material_slot_{slot_index}")
        if name not in name_to_index:
            continue
        selections.append(MaterialSelection(
            index=name_to_index[name],
            name=name,
            contract_item=str(st.session_state.get(f"contract_item_{slot_index}", "")).strip(),
            contract_quantity=str(st.session_state.get(f"contract_quantity_{slot_index}", "")).strip(),
            planned_submission_date=planned_submission_month,
        ))
    return selections


def selection_input_error(selections: list[MaterialSelection]) -> str:
    for position, selection in enumerate(selections, start=1):
        if not selection.contract_item:
            return f"材料 {position} 尚未填寫契約詳細表項次。"
        if not selection.contract_quantity:
            return f"材料 {position} 尚未填寫契約數量。"
    return ""


def planned_submission_month_error(value: str) -> str:
    if value and not re.fullmatch(r"\d{4}\.([1-9]|1[0-2])", value):
        return "預定送審年月格式必須是 YYYY.M，例如 2026.8。"
    return ""


def generation_fingerprint(
    project_name: str,
    selections: list[MaterialSelection],
) -> tuple[str, tuple[tuple[int, str, str, str, str], ...]]:
    return (
        project_name,
        tuple(
            (
                selection.index,
                selection.name,
                selection.contract_item,
                selection.contract_quantity,
                selection.planned_submission_date,
            )
            for selection in selections
        ),
    )


def render_material_form(materials: tuple[str, ...]) -> None:
    project_column, month_column = st.columns([2, 1])
    with project_column:
        st.text_input("工程名稱", placeholder="請輸入工程名稱", key="project_name")
    with month_column:
        st.text_input(
            "預定送審年月（全部材料共用、選填）",
            placeholder="例如 2026.8",
            key="planned_submission_month",
            max_chars=7,
        )
    st.markdown("## 材料選擇表單")
    if not materials:
        st.info("請先從左側匯入四份 Word 檔案，讀取完成後會顯示 8 個材料選單。")
        slot_count = 0
    else:
        slot_count = max(8, int(st.session_state.get("material_slot_count", 8)))
        slot_count = min(slot_count, len(materials))
        st.session_state["material_slot_count"] = slot_count
        st.caption(
            f"已擷取 {len(materials)} 項可選材料；選單依序編號，目前顯示 {slot_count} 個材料欄位。"
        )
        sorted_materials = sorted(materials, key=material_sort_key)
        material_labels = {
            name: f"{position}.{name}"
            for position, name in enumerate(sorted_materials, start=1)
        }
        current_selected = selected_material_names(materials, slot_count)
        for index in range(slot_count):
            key = f"material_slot_{index}"
            current = st.session_state.get(key)
            unavailable = set(current_selected)
            if current:
                unavailable.discard(current)
            options = [""] + [name for name in sorted_materials if name not in unavailable]
            columns = st.columns([2.4, 1.05, 1.05])
            with columns[0]:
                selected_name = st.selectbox(
                    f"材料 {index + 1}",
                    options,
                    key=key,
                    format_func=lambda value, labels=material_labels: labels.get(value, "請選擇材料"),
                )
            identity_key = f"material_slot_identity_{index}"
            if selected_name != st.session_state.get(identity_key):
                for field in ("contract_item", "contract_quantity"):
                    st.session_state.pop(f"{field}_{index}", None)
                st.session_state[identity_key] = selected_name
            with columns[1]:
                st.text_input(
                    "契約詳細表項次",
                    key=f"contract_item_{index}",
                    disabled=not selected_name,
                )
            with columns[2]:
                st.text_input(
                    "契約數量",
                    key=f"contract_quantity_{index}",
                    disabled=not selected_name,
                    placeholder="例如 66CM2、555kg",
                )
        st.button(
            "＋ 新增材料",
            on_click=add_material_slot,
            disabled=slot_count >= len(materials),
        )

    planned_month = st.session_state.get("planned_submission_month", "").strip()
    chosen = material_selections(materials, slot_count, planned_month)
    st.markdown("## 本次產出材料")
    if chosen:
        st.write("、".join(
            f"{position}.{selection.name}"
            for position, selection in enumerate(chosen, start=1)
        ))
    else:
        st.caption("尚未選擇材料")

    _, button_column = st.columns([4, 1])
    with button_column:
        if st.button("完成", type="primary", use_container_width=True, disabled=not chosen):
            project_name = st.session_state.get("project_name", "").strip()
            if not project_name:
                st.error("請先輸入工程名稱。")
            elif month_error := planned_submission_month_error(planned_month):
                st.error(month_error)
            elif input_error := selection_input_error(chosen):
                st.error(input_error)
            else:
                try:
                    stored = st.session_state.get("imported_sources", {})
                    sources = {
                        DocumentKind[kind_name]: UploadedSource(upload_name, content)
                        for kind_name, (upload_name, content) in stored.items()
                    }
                    st.session_state["generated_docx"] = generate_four_form_docx(
                        sources, project_name, chosen,
                    )
                    st.session_state["generated_fingerprint"] = generation_fingerprint(
                        project_name,
                        chosen,
                    )
                except (FourFormGenerationError, KeyError, TypeError, ValueError) as exc:
                    st.session_state.pop("generated_docx", None)
                    st.session_state.pop("generated_fingerprint", None)
                    st.error(f"Word 產生失敗：{exc}")

        project_name = st.session_state.get("project_name", "").strip()
        fingerprint = generation_fingerprint(project_name, chosen)
        if st.session_state.get("generated_docx") and st.session_state.get("generated_fingerprint") == fingerprint:
            safe_name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", project_name).strip(" .")
            st.download_button(
                "下載四表 Word",
                data=st.session_state["generated_docx"],
                file_name=f"{safe_name or '工程材料'}_四表.docx",
                mime=DOCX_MIME,
                use_container_width=True,
                on_click="ignore",
            )


def main() -> None:
    st.set_page_config(page_title=APP_NAME, layout="wide", initial_sidebar_state="expanded")
    apply_theme()
    settings = load_settings(secret_mapping())
    if not require_login(settings):
        return
    materials = render_word_import(settings)
    render_material_form(materials)


def create_password_hash_from_terminal() -> None:
    first = getpass.getpass("請輸入新密碼：")
    second = getpass.getpass("請再次輸入新密碼：")
    if first != second:
        raise SystemExit("兩次密碼不一致。")
    print(hash_password(first))


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--hash-password":
        create_password_hash_from_terminal()
    else:
        main()
